# Copyright 2026 Google LLC
#
# Licensed under the Apache License, Version 2.0 (the "License");
# you may not use this file except in compliance with the License.
# You may obtain a copy of the License at
#
#     https://www.apache.org/licenses/LICENSE-2.0
#
# Unless required by applicable law or agreed to in writing, software
# distributed under the License is distributed on an "AS IS" BASIS,
# WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
# See the License for the specific language governing permissions and
# limitations under the License.

"""DOCX -> Markdown parser. Converts to PDF first, then uses Gemini."""

from __future__ import annotations

from docx import Document

from src.document_preprocessing.parser.convert import (
    parse_via_pdf as parse,  # noqa: F401
)


def quick_raw_text(file_path: str) -> str:
    """Cheap raw-text extraction (no LibreOffice, no LLM).

    Walks the document's paragraphs and table cells via python-docx.
    Skips formatting and embedded images — just enough text for the
    relevance classifier to judge.
    """
    doc = Document(file_path)
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)
    return "\n".join(parts)
